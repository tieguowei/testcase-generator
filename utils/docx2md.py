#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Word文档转Markdown脚本
支持多种元素：标题、文本、图片、表格、列表等
"""

import os
import re
import base64
from pathlib import Path
from typing import List, Optional, Dict, Any
import argparse
import logging

try:
    from docx import Document
    from docx.document import Document as _Document
    from docx.oxml.text.paragraph import CT_P
    from docx.oxml.table import CT_Tbl
    from docx.table import _Cell, Table
    from docx.text.paragraph import Paragraph
    from docx.shared import RGBColor
    from docx.enum.style import WD_STYLE_TYPE
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    import docx.oxml.ns as ns
except ImportError:
    print("请安装必要的依赖包:")
    print("pip install python-docx")
    exit(1)

# 配置日志
logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(levelname)s - %(message)s')
logger = logging.getLogger(__name__)

class WordToMarkdownConverter:
    """Word文档转Markdown转换器"""
    
    def __init__(self, 
                 output_dir: str = "output",
                 # extract_images: bool = True,
                 preserve_formatting: bool = True):
        """
        初始化转换器
        
        Args:
            output_dir: 输出目录
            # extract_images: 是否提取图片
            preserve_formatting: 是否保留格式
        """
        self.output_dir = Path(output_dir)
        self.output_dir.mkdir(exist_ok=True)
        # self.extract_images = extract_images
        self.preserve_formatting = preserve_formatting
        # self.image_counter = 0
        self.markdown_content = []
        # self.image_map = {}  # 存储图片ID到文件名的映射
        # self.image_reference_counter = 0  # 用于追踪已引用的图片数量
        
    def convert(self, docx_path: str, output_name: Optional[str] = None) -> str:
        """
        转换Word文档为Markdown
        
        Args:
            docx_path: Word文档路径
            output_name: 输出文件名（可选）
            
        Returns:
            生成的Markdown文件路径
        """
        try:
            logger.info(f"开始转换文档: {docx_path}")
            
            # 加载Word文档
            doc = Document(docx_path)

            self.markdown_content = []

            # 处理文档内容
            self._process_document(doc)
            
            # 生成输出文件名
            if not output_name:
                output_name = Path(docx_path).stem + ".md"
            
            output_path = self.output_dir / output_name
            
            # 写入Markdown文件
            with open(output_path, 'w', encoding='utf-8') as f:
                f.write('\n'.join(self.markdown_content))
            
            logger.info(f"转换完成，输出文件: {output_path}")
            return str(output_path)
            
        except Exception as e:
            logger.error(f"转换过程中出错: {str(e)}")
            raise
    
    def _process_document(self, doc: _Document):
        """处理整个文档"""
        if doc is None:
            logger.error("文档对象为空")
            return

        
        try:
            # 遍历文档中的所有元素
            if doc.element is None or doc.element.body is None:
                logger.warning("文档结构异常，尝试使用段落和表格")
                # 尝试直接访问段落和表格
                try:
                    for paragraph in doc.paragraphs:
                        self._process_paragraph(paragraph)
                    for table in doc.tables:
                        self._process_table(table)
                except Exception as e:
                    logger.error(f"处理文档段落和表格时出错: {str(e)}")
                return
            
            for element in doc.element.body:
                try:
                    if isinstance(element, CT_P):
                        # 段落元素
                        paragraph = Paragraph(element, doc)
                        self._process_paragraph(paragraph)
                    elif isinstance(element, CT_Tbl):
                        # 表格元素
                        table = Table(element, doc)
                        self._process_table(table)
                except Exception as e:
                    logger.warning(f"处理文档元素时出错: {str(e)}")
                    continue
                    
        except Exception as e:
            logger.error(f"处理文档时出现严重错误: {str(e)}")
            # 最后的备用方案：只提取纯文本
            try:
                logger.info("尝试提取纯文本作为备用方案")
                for paragraph in doc.paragraphs:
                    if paragraph and paragraph.text:
                        text = paragraph.text.strip()
                        if text:
                            self.markdown_content.append(text)
                            self.markdown_content.append('')
            except Exception as fallback_e:
                logger.error(f"备用文本提取也失败: {str(fallback_e)}")
        
        # 后处理：确保所有提取的图片都被引用
        # self._post_process_images()
    
    def _process_paragraph(self, paragraph: Paragraph):
        """处理段落"""
        if paragraph is None:
            return
            
        try:
            # 先处理段落内容
            # 检查是否为标题
            if self._is_heading(paragraph):
                self._process_heading(paragraph)
            # 检查是否为列表
            elif self._is_list_item(paragraph):
                self._process_list_item(paragraph)
            else:
                # 普通段落
                self._process_normal_paragraph(paragraph)

                
        except Exception as e:
            logger.warning(f"处理段落时出错: {str(e)}")
            # 尝试至少提取纯文本
            try:
                text = paragraph.text if paragraph.text else ""
                if text.strip():
                    self.markdown_content.append(text.strip())
                    self.markdown_content.append('')
            except:
                pass
    

    
    def _is_heading(self, paragraph: Paragraph) -> bool:
        """判断是否为标题"""
        if paragraph.style is None or paragraph.style.name is None:
            return False
        style_name = paragraph.style.name.lower()
        return 'heading' in style_name or style_name.startswith('标题')
    
    def _process_heading(self, paragraph: Paragraph):
        """处理标题"""
        try:
            # 提取标题级别
            level = 1
            if paragraph.style and paragraph.style.name:
                style_name = paragraph.style.name.lower()
                if 'heading' in style_name:
                    # 尝试从样式名中提取级别
                    import re
                    level_match = re.search(r'(\d+)', style_name)
                    if level_match:
                        level = int(level_match.group(1))
                elif '标题' in style_name:
                    # 处理中文标题样式
                    level_match = re.search(r'(\d+)', style_name)
                    if level_match:
                        level = int(level_match.group(1))
            
            # 生成Markdown标题
            text = self._extract_text_with_formatting(paragraph)
            if text.strip():
                heading_text = '#' * level + ' ' + text.strip()
                self.markdown_content.append(heading_text)
                self.markdown_content.append('')
                
                # 在标题后尝试插入图片（更高优先级）
                # self._try_insert_next_image(force_insert=True)
                
        except Exception as e:
            logger.warning(f"处理标题时出错: {str(e)}")
            # 备用方案：至少提取文本
            try:
                text = paragraph.text if paragraph.text else ""
                if text.strip():
                    self.markdown_content.append(f"# {text.strip()}")
                    self.markdown_content.append('')
            except:
                pass
    
    def _is_list_item(self, paragraph: Paragraph) -> bool:
        """判断是否为列表项"""
        if paragraph is None:
            return False
            
        try:
            # 检查段落的编号样式
            if paragraph._element is not None:
                pPr = paragraph._element.pPr
                if pPr is not None:
                    numPr = pPr.numPr
                    if numPr is not None:
                        return True
        except Exception as e:
            logger.debug(f"检查段落编号样式时出错: {str(e)}")
        
        try:
            # 检查文本开头是否为列表标记
            text = paragraph.text
            if text:
                text = text.strip()
                if text:
                    # 无序列表标记
                    if text.startswith(('•', '·', '○', '■', '▪', '-', '*')):
                        return True
                    # 有序列表标记
                    if re.match(r'^\d+\.', text) or re.match(r'^[a-zA-Z]\.', text):
                        return True
        except Exception as e:
            logger.debug(f"检查段落文本时出错: {str(e)}")
        
        return False
    
    def _process_list_item(self, paragraph: Paragraph):
        """处理列表项"""
        text = self._extract_text_with_formatting(paragraph)
        if not text.strip():
            return
        
        # 检查是否为有序列表
        is_ordered = False
        pPr = paragraph._element.pPr
        if pPr is not None:
            numPr = pPr.numPr
            if numPr is not None:
                # 这里可以进一步判断列表类型
                is_ordered = True
        
        # 简单的文本匹配判断
        if re.match(r'^\d+\.', text.strip()):
            is_ordered = True
        
        # 移除原有的列表标记
        cleaned_text = re.sub(r'^[•·○■▪\-\*]\s*', '', text.strip())
        cleaned_text = re.sub(r'^\d+\.\s*', '', cleaned_text)
        cleaned_text = re.sub(r'^[a-zA-Z]\.\s*', '', cleaned_text)
        
        # 生成Markdown列表项
        if is_ordered:
            self.markdown_content.append(f"1. {cleaned_text}")
        else:
            self.markdown_content.append(f"- {cleaned_text}")
    
    def _process_normal_paragraph(self, paragraph: Paragraph):
        """处理普通段落"""
        text = self._extract_text_with_formatting(paragraph)
        
        if text.strip():
            self.markdown_content.append(text)
            self.markdown_content.append('')  # 添加空行
    
    def _extract_text_with_formatting(self, paragraph: Paragraph) -> str:
        """提取段落文本并保留格式"""
        if paragraph is None:
            return ""
        
        result = []
        
        try:
            for run in paragraph.runs:
                if run is None:
                    continue
                    
                text = run.text
                if not text:
                    continue
                
                if self.preserve_formatting:
                    try:
                        # 应用格式
                        if run.bold:
                            text = f"**{text}**"
                        if run.italic:
                            text = f"*{text}*"
                        if run.underline:
                            text = f"<u>{text}</u>"
                        
                        # 处理删除线
                        if hasattr(run.font, 'strike') and run.font.strike:
                            text = f"~~{text}~~"
                        
                        # 处理上标和下标
                        if run.font and hasattr(run.font, 'superscript') and run.font.superscript:
                            text = f"<sup>{text}</sup>"
                        elif run.font and hasattr(run.font, 'subscript') and run.font.subscript:
                            text = f"<sub>{text}</sub>"
                    except Exception as e:
                        logger.debug(f"处理文本格式时出错: {str(e)}")
                
                result.append(text)
        except Exception as e:
            logger.warning(f"提取段落文本时出错: {str(e)}")
            # 如果格式化处理失败，至少返回纯文本
            try:
                return paragraph.text if paragraph.text else ""
            except:
                return ""
        
        return ''.join(result)
    
    def _process_table(self, table: Table):
        """处理表格"""
        if table is None:
            return
            
        try:
            if not table.rows:
                return
            
            markdown_table = []
            
            # 处理表头
            if table.rows:
                try:
                    header_row = table.rows[0]
                    headers = []
                    for cell in header_row.cells:
                        cell_text = self._extract_cell_text(cell)
                        headers.append(cell_text.strip())
                    
                    if headers:  # 确保有表头内容
                        markdown_table.append('| ' + ' | '.join(headers) + ' |')
                        markdown_table.append('| ' + ' | '.join(['---'] * len(headers)) + ' |')
                except Exception as e:
                    logger.warning(f"处理表头时出错: {str(e)}")
                    return
            
            # 处理数据行
            try:
                for row in table.rows[1:]:
                    row_data = []
                    for cell in row.cells:
                        try:
                            cell_text = self._extract_cell_text(cell)
                            # 替换表格中的换行符
                            cell_text = cell_text.replace('\n', '<br>')
                            row_data.append(cell_text.strip())
                        except Exception as e:
                            logger.debug(f"处理单元格时出错: {str(e)}")
                            row_data.append("")  # 添加空单元格
                    
                    if row_data:  # 确保行有数据
                        markdown_table.append('| ' + ' | '.join(row_data) + ' |')
            except Exception as e:
                logger.warning(f"处理表格行时出错: {str(e)}")
            
            # 添加表格到内容
            if markdown_table:
                self.markdown_content.extend(markdown_table)
                self.markdown_content.append('')  # 添加空行
                
        except Exception as e:
            logger.warning(f"处理表格时出错: {str(e)}")
    
    def _extract_cell_text(self, cell: _Cell) -> str:
        """提取单元格文本"""
        if cell is None:
            return ""
            
        texts = []
        try:
            for paragraph in cell.paragraphs:
                if paragraph is None:
                    continue
                text = self._extract_text_with_formatting(paragraph)
                if text and text.strip():
                    texts.append(text.strip())
        except Exception as e:
            logger.debug(f"提取单元格文本时出错: {str(e)}")
            # 尝试直接获取单元格文本
            try:
                if hasattr(cell, 'text') and cell.text:
                    return cell.text.strip()
            except:
                pass
        
        return ' '.join(texts)
    


def main():
    """主函数"""
    parser = argparse.ArgumentParser(description='Word文档转Markdown工具')
    parser.add_argument('input', help='输入的Word文档路径')
    parser.add_argument('-o', '--output', help='输出文件名（默认为输入文件名，扩展名为.md）')
    parser.add_argument('-d', '--output-dir', default='output', help='输出目录（默认为output）')
    parser.add_argument('--no-formatting', action='store_true', help='不保留格式')
    parser.add_argument('-v', '--verbose', action='store_true', help='详细输出')
    
    args = parser.parse_args()
    
    # 为 -o 参数设置默认值
    if args.output is None:
        # 基于输入文件名生成默认输出文件名
        input_name = os.path.splitext(os.path.basename(args.input))[0]
        args.output = f"{input_name}.md"
    
    if args.verbose:
        logging.getLogger().setLevel(logging.DEBUG)
    
    # 检查输入文件
    if not os.path.exists(args.input):
        print(f"错误: 文件 '{args.input}' 不存在")
        return 1
    
    # 创建转换器
    converter = WordToMarkdownConverter(
        output_dir=args.output_dir,
        # extract_images=not args.no_images,
        preserve_formatting=not args.no_formatting
    )
    
    try:
        # 执行转换
        output_path = converter.convert(args.input, args.output)
        print(f"转换成功! 输出文件: {output_path}")
        return 0
    except Exception as e:
        print(f"转换失败: {str(e)}")
        return 1

if __name__ == '__main__':
    exit(main())